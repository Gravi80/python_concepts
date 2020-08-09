# document_obj -> paragraph_obj -> run_obj
# http://autbor.com/demo.docx
import os
import docx
from docx.document import Document

document: Document = docx.Document(f"{os.getenv('HOME')}/Downloads/demo.docx")
print(document.paragraphs)
print(document.paragraphs[0].text)
print(document.paragraphs[0].runs)  # A new run starts whenever there is a change in the style
print(document.paragraphs[1].runs)
print(document.paragraphs[1].runs[1].bold)
print(document.paragraphs[1].runs[1].text)

print("\n********* Update **********")
print(document.paragraphs[1].runs[3].italic)
print(document.paragraphs[1].runs[3].text)

document.paragraphs[1].runs[3].underline = True
document.paragraphs[1].runs[3].text = 'italic and underlined.'
document.save(f"{os.getenv('HOME')}/Downloads/demo2.docx")

print(document.paragraphs[1].style)
document.paragraphs[1].style = 'Title'
document.save(f"{os.getenv('HOME')}/Downloads/demo3.docx")

print("\n********* New word document **********")
new_doc: Document = docx.Document()
new_doc.add_paragraph('Hello this is a paragraph')
new_doc.add_paragraph('Another paragraph')
new_doc.save(f"{os.getenv('HOME')}/Downloads/demo4.docx")

new_doc.paragraphs[0].add_run("This is a new run")
new_doc.paragraphs[0].runs[1].bold = True
new_doc.save(f"{os.getenv('HOME')}/Downloads/demo5.docx")
